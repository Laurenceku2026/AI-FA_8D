"""
AI-FA 智能故障分析系统
AI-powered Failure Analysis & 8D Report Generation

版本: 3.1.0
更新: 
- 集成 TechLife Portal 计费逻辑
- 侧边栏显示用户名和实时剩余次数
- 三个按钮（分析、FA报告、8D报告）各自扣费
- 分析完成后才启用报告按钮
- 免费版次数不足时按钮禁用
- 专业版无限使用
"""

import streamlit as st
import pandas as pd
import numpy as np
import json
import uuid
import re
import io
import requests
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from PIL import Image
import plotly.graph_objects as go
import plotly.express as px
import matplotlib.pyplot as plt
from matplotlib.font_manager import fontManager

# ==================== 配置（从 Streamlit Secrets 读取）====================

def get_secret(key: str, default: str = "") -> str:
    """安全获取Secret配置"""
    try:
        if hasattr(st, 'secrets') and key in st.secrets:
            return st.secrets[key]
    except:
        pass
    return default

# DeepSeek API 配置
DEEPSEEK_API_KEY = get_secret("DEEPSEEK_API_KEY")
DEEPSEEK_BASE_URL = get_secret("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
DEEPSEEK_MODEL = get_secret("DEEPSEEK_MODEL", "deepseek-chat")

# Supabase 配置
SUPABASE_URL = get_secret("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = get_secret("SUPABASE_SERVICE_ROLE_KEY")

# 管理员配置
ADMIN_USERNAME = "Laurence_ku"
ADMIN_PASSWORD = "Ku_product$2026"

# 专业术语白名单（这些词在任何语言下都保留原样）
TECHNICAL_TERMS = [
    "LED", "PCB", "PCBA", "SMD", "COB", "IC", "MCU", "MOSFET", "Diode", "Capacitor", "Resistor",
    "DMX", "PWM", "DALI", "SPI", "I2C", "UART", "RS485",
    "UL", "CE", "RoHS", "IP65", "IP67", "IP68", "IK10",
    "RGB", "RGBW", "CCT", "CRI", "ESD", "EMI", "EMC",
    "AC", "DC", "VAC", "VDC", "mA", "A", "V", "W", "lm", "K", "nm"
]

# 页面配置
st.set_page_config(
    page_title="AI-FA 智能故障分析系统",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ==================== 接收门户参数 ====================

query_params = st.query_params


def set_app_language(lang: str):
    """同步 session 与 URL 语言，避免 query_params 在 rerun 时覆盖用户选择。"""
    if lang not in ("zh", "en"):
        lang = "zh"
    st.session_state.lang = lang
    st.query_params["lang"] = lang


if "user_id" in query_params:
    # 获取 user_id
    user_id_val = query_params["user_id"]
    if isinstance(user_id_val, list):
        st.session_state.user_id = user_id_val[0]
    else:
        st.session_state.user_id = user_id_val
    
    # 获取 email
    email_val = query_params.get("email", "")
    if isinstance(email_val, list):
        st.session_state.user_email = email_val[0] if email_val else ""
    else:
        st.session_state.user_email = email_val
    
    # 从邮箱提取用户名
    if st.session_state.user_email and "@" in st.session_state.user_email:
        st.session_state.username = st.session_state.user_email.split('@')[0]
    else:
        st.session_state.username = "User"
    
    # 设置语言（仅首次进入时从 URL 读取）
    if "lang" not in st.session_state:
        if "lang" in query_params:
            lang_val = query_params["lang"]
            if isinstance(lang_val, list):
                lang_val = lang_val[0]
            set_app_language(lang_val if lang_val in ["zh", "en"] else "zh")
        else:
            set_app_language("zh")
    
    # 接收剩余次数（仅用于初始显示）
    if "trials_left" in query_params:
        trials_val = query_params["trials_left"]
        if isinstance(trials_val, list):
            trials_val = trials_val[0]
        st.session_state.trials_left = int(trials_val)
else:
    st.warning("请从 TechLife Suite 门户登录后访问")
    st.stop()


# ==================== Supabase 配置 ====================

HEADERS = {
    "apikey": SUPABASE_SERVICE_ROLE_KEY,
    "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
    "Content-Type": "application/json"
}


def supabase_get(table: str, user_id: str = None):
    """GET 请求"""
    url = f"{SUPABASE_URL}/rest/v1/{table}"
    if user_id:
        url += f"?id=eq.{user_id}"
    response = requests.get(url, headers=HEADERS)
    return response


def supabase_patch(table: str, user_id: str, data: dict):
    """PATCH 请求（更新）"""
    url = f"{SUPABASE_URL}/rest/v1/{table}?id=eq.{user_id}"
    response = requests.patch(url, headers=HEADERS, json=data)
    return response


def supabase_post(table: str, data: dict):
    """POST 请求"""
    url = f"{SUPABASE_URL}/rest/v1/{table}"
    response = requests.post(url, headers=HEADERS, json=data)
    return response


def get_user_remaining_trials(user_id: str) -> tuple:
    """
    从数据库实时获取剩余次数和订阅类型
    返回: (remaining, tier, expires_at, error_msg)
    - remaining: -1 表示无限（专业版），>=0 表示剩余次数
    - tier: "free" 或 "pro"
    """
    try:
        response = supabase_get("profiles", user_id)
        if response.status_code == 200 and response.json():
            data = response.json()[0]
            tier = data.get("subscription_tier", "free")
            remaining = data.get("free_trials_remaining", 30)
            expires_at = data.get("subscription_expires_at")
            
            if tier == "pro":
                return -1, "pro", expires_at, ""
            return remaining, "free", expires_at, ""
    except Exception as e:
        return None, None, None, f"查询失败: {str(e)}"
    
    return 30, "free", None, ""


def consume_trial(user_id: str, app_name: str, action_name: str) -> tuple:
    """
    消耗一次免费次数
    返回: (是否成功, 剩余次数, 错误信息)
    """
    try:
        # 获取当前用户状态
        resp = supabase_get("profiles", user_id)
        if resp.status_code != 200 or not resp.json():
            return False, 0, "用户不存在"
        
        current = resp.json()[0].get("free_trials_remaining", 30)
        tier = resp.json()[0].get("subscription_tier", "free")
        
        # 专业版无限使用
        if tier == "pro":
            return True, -1, ""
        
        if current <= 0:
            return False, 0, f"免费次数已用完（共30次），请联系管理员升级。\n\n当前剩余: {current} 次"
        
        # 更新剩余次数
        patch_resp = supabase_patch("profiles", user_id, {"free_trials_remaining": current - 1})
        
        if patch_resp.status_code not in [200, 204]:
            return False, 0, f"更新失败: {patch_resp.text}"
        
        # 记录使用日志
        supabase_post("usage_logs", {
            "user_id": user_id,
            "app_name": app_name,
            "action_name": action_name,
            "analysis_count": 1,
            "used_at": datetime.now().isoformat()
        })
        
        return True, current - 1, ""
        
    except Exception as e:
        return False, 0, f"计数失败: {str(e)}"


# ==================== 侧边栏（显示用户信息和实时剩余次数）===================

def render_sidebar_user_info():
    """渲染侧边栏用户信息和剩余次数（实时查询）"""
    with st.sidebar:
        # 显示用户名
        st.markdown(f"### 👤 {st.session_state.username}")
        
        # 实时查询剩余次数
        remaining, tier, expires_at, error = get_user_remaining_trials(st.session_state.user_id)
        
        if error:
            st.error(error)
        else:
            if tier == "pro":
                st.info("🎫 剩余免费次数: ∞ (专业版)")
                if expires_at:
                    st.caption(f"📅 到期: {expires_at[:10]}")
            else:
                st.info(f"🎫 剩余免费次数: {remaining}")
        
        st.markdown("---")


# ==================== 数据模型（双语版本）====================

@dataclass
class FiveWhyItem:
    """5-Why推理项（双语）"""
    level: int
    question_en: str
    question_zh: str
    answer_en: str
    answer_zh: str
    confidence: float
    verification_method: str
    
    def to_dict(self) -> dict:
        return {
            "level": self.level,
            "question_en": self.question_en,
            "question_zh": self.question_zh,
            "answer_en": self.answer_en,
            "answer_zh": self.answer_zh,
            "confidence": self.confidence,
            "verification_method": self.verification_method
        }


@dataclass
class FishboneAnalysis:
    """鱼骨图分析结果（双语）"""
    man_en: List[str] = field(default_factory=list)
    man_zh: List[str] = field(default_factory=list)
    machine_en: List[str] = field(default_factory=list)
    machine_zh: List[str] = field(default_factory=list)
    material_en: List[str] = field(default_factory=list)
    material_zh: List[str] = field(default_factory=list)
    method_en: List[str] = field(default_factory=list)
    method_zh: List[str] = field(default_factory=list)
    environment_en: List[str] = field(default_factory=list)
    environment_zh: List[str] = field(default_factory=list)
    measurement_en: List[str] = field(default_factory=list)
    measurement_zh: List[str] = field(default_factory=list)
    
    def get_causes(self, lang: str, category: str) -> List[str]:
        """获取指定语言和类别的原因列表"""
        if lang == "zh":
            mapping = {
                "Man": self.man_zh, "Machine": self.machine_zh,
                "Material": self.material_zh, "Method": self.method_zh,
                "Environment": self.environment_zh, "Measurement": self.measurement_zh
            }
        else:
            mapping = {
                "Man": self.man_en, "Machine": self.machine_en,
                "Material": self.material_en, "Method": self.method_en,
                "Environment": self.environment_en, "Measurement": self.measurement_en
            }
        return mapping.get(category, [])
    
    def to_dict(self, lang: str = "zh") -> dict:
        """输出指定语言的字典"""
        if lang == "zh":
            return {
                "Man": self.man_zh, "Machine": self.machine_zh,
                "Material": self.material_zh, "Method": self.method_zh,
                "Environment": self.environment_zh, "Measurement": self.measurement_zh
            }
        else:
            return {
                "Man": self.man_en, "Machine": self.machine_en,
                "Material": self.material_en, "Method": self.method_en,
                "Environment": self.environment_en, "Measurement": self.measurement_en
            }


@dataclass
class FailureAnalysisResult:
    """故障分析完整结果（双语）"""
    case_id: str
    project_name: str
    product_name: str
    symptom: str                      # 原始用户输入（保留）
    symptom_en: str                   # 英文版本
    installation: str                 # 原始用户输入
    installation_en: str              # 英文版本
    temperature: str
    failure_stage: int
    five_why: List[FiveWhyItem]
    fishbone: FishboneAnalysis
    root_cause_en: str
    root_cause_zh: str
    root_cause_confidence: float
    interim_actions_en: List[str]
    interim_actions_zh: List[str]
    permanent_actions_en: List[str]
    permanent_actions_zh: List[str]
    preventive_actions_en: List[str]
    preventive_actions_zh: List[str]
    internal_cases_used: int
    external_sources_used: int
    spc_analysis: Optional[dict] = None
    association_rules: Optional[List[dict]] = None
    analyst_name: str = ""
    analyst_title: str = ""
    fishbone_image: Optional[bytes] = None


# ==================== 多语言文本 ====================

TEXTS = {
    "zh": {
        "app_title": "AI-FA 智能故障分析系统",
        "app_subtitle": "AI驱动的失效分析与8D报告生成",
        "lang_zh": "中文",
        "lang_en": "English",
        "sidebar_about": "关于系统",
        "sidebar_principle": "原理：基于DeepSeek大语言模型 + 5-Why推理引擎，自动进行失效等级分类、鱼骨图分析，结合历史案例库定位根本原因。",
        "sidebar_usage": "使用方法：\n1. 输入产品名称和故障现象\n2. 可选：上传图片、时序数据\n3. 点击「开始AI深度故障分析」\n4. 选择生成FA报告或8D报告\n5. 下载Word文档",
        "analyst_name": "分析人姓名",
        "analyst_name_ph": "请输入姓名",
        "analyst_title": "分析人头衔（可选）",
        "analyst_title_ph": "例如：质量总监",
        "project_name_label": "项目名称（可选）",
        "project_name_ph": "例如：启德体育场项目",
        "db_status": "数据库状态",
        "db_connected": "已连接",
        "db_disconnected": "未连接",
        "contact": "联系",
        "contact_email": "电邮: Techlife2027@gmail.com",
        
        "basic_info": "故障基本信息",
        "product_name": "产品名称",
        "product_name_ph": "例如：Apollo Dot L RGBW",
        "symptom": "故障现象",
        "symptom_ph": "例如：暴雨后灯具闪烁、不亮、有烧焦痕迹...",
        "installation": "安装位置/方向",
        "installation_ph": "例如：上层迎光面、黑色金属表面",
        "failure_date": "故障发生时间",
        "batch_no": "批次/序列号",
        "site_temp": "现场温度(可选)",
        "site_temp_ph": "例如：70°C",
        "image_section": "故障图片",
        "image_upload_hint": "点击上传图片",
        "image_paste_hint": "或从剪贴板粘贴",
        "image_paste_btn": "粘贴图片",
        "image_paste_success": "图片已粘贴",
        "image_paste_error": "剪贴板无图片，请先截图",
        "no_images": "暂无图片",
        "timeseries_section": "时序数据分析",
        "timeseries_checkbox": "启用时序数据分析（SPC控制图）",
        "timeseries_input_method": "数据输入方式",
        "timeseries_paste": "直接粘贴数据",
        "timeseries_upload": "上传Excel/CSV文件",
        "timeseries_paste_placeholder": "请输入时序数据，格式：\n日期,生产数量,故障数量\n2024-01-01,1000,5\n2024-02-01,1100,8",
        "timeseries_file_hint": "支持 .xlsx, .xls, .csv 格式",
        "download_template": "下载数据模板",
        
        "advanced_options": "高级分析选项",
        "web_search": "联网搜索行业案例",
        "rule_mining": "关联规则挖掘",
        "spc": "时序分析(SPC控制图)",
        "gen_8d": "生成8D报告",
        
        "analyze_btn": "开始AI深度故障分析",
        "generate_fa_btn": "生成FA报告",
        "generate_8d_btn": "生成8D报告",
        "download_word": "下载Word报告",
        "clear_btn": "清除结果",
        
        "five_why_title": "5-Why 根因分析",
        "fishbone_title": "鱼骨图分析",
        "root_cause_title": "根因结论",
        "interim_actions": "临时措施",
        "permanent_actions": "永久措施",
        "preventive_actions": "预防措施",
        "confidence": "置信度",
        "stage_label": "失效等级",
        "spc_title": "SPC控制图分析",
        "rules_title": "关联规则挖掘",
        
        "analyzing": "AI正在分析中，请稍候...",
        "success": "分析完成！",
        "error": "分析失败，请重试",
        "fill_required": "请填写产品名称和故障现象",
        "api_error": "API配置错误，请检查Streamlit Secrets",
        
        "stage_0": "正常",
        "stage_1": "轻微异常",
        "stage_2": "中度异常",
        "stage_3": "严重故障",
                # 失效等级表格（中文）
        "stage_table_title": "等级定义参考",
        "stage_table_header_level": "等级",
        "stage_table_header_name": "名称",
        "stage_table_header_characteristics": "典型特征",
        "stage_table_header_action": "建议行动",
        "stage_0_char": "功能正常，无异常",
        "stage_0_action": "无需处理",
        "stage_1_char": "闪烁、弱光、色偏",
        "stage_1_action": "监控观察",
        "stage_2_char": "烧焦痕迹、变形",
        "stage_2_action": "立即更换",
        "stage_3_char": "短路、冒烟、起火",
        "stage_3_action": "紧急停机",
        "report_preview": "报告预览",
        "fault_photos": "故障照片",
        "analyst": "分析人",
        "analysis_date": "分析日期",
        "project_name_header": "项目名称",
        "product_name_header": "产品名称",
        "fault_summary": "故障简述",
        "report_info": "报告信息",
        "report_date": "报告日期",
        
        "what": "故障现象",
        "where": "发现位置",
        "when": "发生时间",
        "who": "发现人",
        "why": "初步判断",
        "how": "发生过程",
        "how_many": "影响数量",
        
        "admin_settings": "管理员设置",
        "admin_login": "管理员验证",
        "username": "用户名",
        "password": "密码",
        "login_btn": "登录",
        "logged_in": "管理员已登录",
        "llm_status": "大模型状态",
        "llm_configured": "DeepSeek API 已配置",
        "llm_not_configured": "DeepSeek API 未配置",
        "neo4j_status": "Neo4j 知识图谱",
        "neo4j_connected": "已连接",
        "neo4j_disconnected": "未连接",
        "knowledge_base_title": "知识库管理（双语）",
        "category": "选择分类",
        "add_entry": "添加条目",
        "entry_content": "经验教训内容",
        "entry_placeholder": "输入经验教训，系统会自动翻译存储双语",
        "no_entries": "暂无条目",
        "export_kb": "导出知识库",
        "import_kb": "导入知识库（Excel）",
        "import_success": "导入成功！共导入 {count} 条记录",

        "team_role": "角色",
        "team_responsibility": "职责",
        "team_leader": "团队负责人",
        "team_leader_resp": "总体协调",
        "design_engineer": "设计工程师",
        "design_engineer_resp": "技术分析",
        "quality_engineer": "质量工程师",
        "quality_engineer_resp": "质量验证",
        "discovery_location": "发现位置",
        "discovery_person": "发现人",
        "preliminary_judgment": "初步判断",
        "occurrence_process": "发生过程",
        "affected_quantity": "影响数量",
        "installation_site": "安装现场",
        "field_maintenance_team": "现场维护团队",
        "under_investigation": "调查中",
        "failure_during_operation": "运行过程中发生故障",
        "fishbone_analysis_title": "鱼骨图分析",
        "five_why_analysis_title": "5-Why 分析",
        "verified_root_cause_title": "验证后的根本原因",
        "verification_item": "项目",
        "verification_method": "方法",
        "verification_criteria": "标准",
        "function_test": "功能",
        "function_test_method": "实际测试",
        "function_test_criteria": "正常运行",
        "durability_test": "耐久性",
        "durability_test_method": "加速老化",
        "durability_test_criteria": "满足设计寿命",
        "d8_recognition_1": "根本原因已确认",
        "d8_recognition_2": "改进措施已定义",
        "d8_recognition_3": "经验教训已加入知识库",
        
        "establish_team": "建立团队",
        "problem_description": "问题描述",
        "root_cause_analysis": "根本原因分析",
        "effectiveness_verification": "效果验证",
        "team_recognition": "总结表彰",
        "table_format": "表格形式",
        "detailed_description": "详细说明",
        "question_label": "问题",
        "answer_label": "答案",
        
        "trials_insufficient": "免费次数已用完，请联系管理员升级到专业版",
        "click_to_upgrade": "升级专业版",
    },
    "en": {
        "app_title": "AI-FA Intelligent Failure Analysis System",
        "app_subtitle": "AI-powered Failure Analysis & 8D Report Generation",
        "lang_zh": "Chinese",
        "lang_en": "English",
        "sidebar_about": "About",
        "sidebar_principle": "Principle: DeepSeek LLM + 5-Why reasoning engine for failure stage classification, fishbone analysis, and root cause identification.",
        "sidebar_usage": "How to use:\n1. Enter product name and failure symptom\n2. Optional: Upload images, time series data\n3. Click 'Start AI Deep Failure Analysis'\n4. Select FA report or 8D report\n5. Download Word document",
        "analyst_name": "Analyst Name",
        "analyst_name_ph": "Enter name",
        "analyst_title": "Title (Optional)",
        "analyst_title_ph": "e.g., Quality Director",
        "project_name_label": "Project Name (Optional)",
        "project_name_ph": "e.g., Kai Tak Stadium Project",
        "db_status": "Database Status",
        "db_connected": "Connected",
        "db_disconnected": "Disconnected",
        "contact": "Contact",
        "contact_email": "Email: Techlife2027@gmail.com",
        
        "basic_info": "Basic Failure Information",
        "product_name": "Product Name",
        "product_name_ph": "e.g., Apollo Dot L RGBW",
        "symptom": "Failure Symptom",
        "symptom_ph": "e.g., flickering, no light output, burning marks after heavy rain...",
        "installation": "Installation Position/Orientation",
        "installation_ph": "e.g., upper facade, black metal surface",
        "failure_date": "Failure Date",
        "batch_no": "Batch/Serial No.",
        "site_temp": "Site Temperature (Optional)",
        "site_temp_ph": "e.g., 70°C",
        "image_section": "Failure Images",
        "image_upload_hint": "Click to upload",
        "image_paste_hint": "or paste from clipboard",
        "image_paste_btn": "Paste Image",
        "image_paste_success": "Image pasted",
        "image_paste_error": "No image in clipboard",
        "no_images": "No images",
        "timeseries_section": "Time Series Analysis",
        "timeseries_checkbox": "Enable Time Series Analysis (SPC Chart)",
        "timeseries_input_method": "Data Input Method",
        "timeseries_paste": "Paste Data",
        "timeseries_upload": "Upload Excel/CSV",
        "timeseries_paste_placeholder": "Paste time series data, format:\ndate,production_qty,failure_qty\n2024-01-01,1000,5\n2024-02-01,1100,8",
        "timeseries_file_hint": "Supports .xlsx, .xls, .csv",
        "download_template": "Download Template",
        
        "advanced_options": "Advanced Analysis Options",
        "web_search": "Web search for industry cases",
        "rule_mining": "Association rule mining",
        "spc": "Time series analysis (SPC)",
        "gen_8d": "Generate 8D report",
        
        "analyze_btn": "Start AI Deep Failure Analysis",
        "generate_fa_btn": "Generate FA Report",
        "generate_8d_btn": "Generate 8D Report",
        "download_word": "Download Word Report",
        "clear_btn": "Clear Results",
        
        "five_why_title": "5-Why Root Cause Analysis",
        "fishbone_title": "Fishbone Diagram",
        "root_cause_title": "Root Cause",
        "interim_actions": "Interim Actions",
        "permanent_actions": "Permanent Actions",
        "preventive_actions": "Preventive Actions",
        "confidence": "Confidence",
        "stage_label": "Failure Stage",
        "spc_title": "SPC Control Chart Analysis",
        "rules_title": "Association Rule Mining",
        
        "analyzing": "AI is analyzing, please wait...",
        "success": "Analysis completed!",
        "error": "Analysis failed, please retry",
        "fill_required": "Please fill in product name and symptom",
        "api_error": "API configuration error, please check Streamlit Secrets",
        
        "stage_0": "Normal",
        "stage_1": "Minor Anomaly",
        "stage_2": "Moderate Anomaly",
        "stage_3": "Critical Failure",
                # 失效等级表格（英文）
        "stage_table_title": "Stage Definition Reference",
        "stage_table_header_level": "Level",
        "stage_table_header_name": "Name",
        "stage_table_header_characteristics": "Typical Characteristics",
        "stage_table_header_action": "Recommended Action",
        "stage_0_char": "Normal function, no abnormality",
        "stage_0_action": "No action needed",
        "stage_1_char": "Flickering, dim light, color shift",
        "stage_1_action": "Monitor",
        "stage_2_char": "Burning marks, deformation",
        "stage_2_action": "Replace immediately",
        "stage_3_char": "Short circuit, smoke, fire",
        "stage_3_action": "Emergency shutdown",
        "report_preview": "Report Preview",
        "fault_photos": "Fault Photos",
        "analyst": "Analyst",
        "analysis_date": "Analysis Date",
        "project_name_header": "Project Name",
        "product_name_header": "Product Name",
        "fault_summary": "Fault Summary",
        "report_info": "Report Information",
        "report_date": "Report Date",
        
        "what": "What",
        "where": "Where",
        "when": "When",
        "who": "Who",
        "why": "Why (Preliminary)",
        "how": "How",
        "how_many": "How Many",
        
        "admin_settings": "Admin Settings",
        "admin_login": "Admin Verification",
        "username": "Username",
        "password": "Password",
        "login_btn": "Login",
        "logged_in": "Admin logged in",
        "llm_status": "LLM Status",
        "llm_configured": "DeepSeek API Configured",
        "llm_not_configured": "DeepSeek API Not Configured",
        "neo4j_status": "Neo4j Knowledge Graph",
        "neo4j_connected": "Connected",
        "neo4j_disconnected": "Not Connected",
        "knowledge_base_title": "Knowledge Base (Bilingual)",
        "category": "Select Category",
        "add_entry": "Add Entry",
        "entry_content": "Lesson Content",
        "entry_placeholder": "Enter lesson, auto-translated to bilingual",
        "no_entries": "No entries",
        "export_kb": "Export Knowledge Base",
        "import_kb": "Import Knowledge Base (Excel)",
        "import_success": "Import successful! {count} records imported",
        
        "establish_team": "Establish Team",
        "problem_description": "Problem Description",
        "root_cause_analysis": "Root Cause Analysis",
        "effectiveness_verification": "Effectiveness Verification",
        "team_recognition": "Team Recognition",
        "table_format": "Table Format",
        "detailed_description": "Detailed Description",
        "question_label": "Question",
        "answer_label": "Answer",
        
        "trials_insufficient": "Free trials exhausted. Please upgrade to Pro.",
        "click_to_upgrade": "Upgrade to Pro",
    }
}


def get_text(key: str) -> str:
    """获取当前语言的文本"""
    lang = st.session_state.get("lang", "zh")
    return TEXTS[lang].get(key, key)


# ==================== 工具函数 ====================

def is_chinese(text: str) -> bool:
    """判断是否包含中文字符"""
    return bool(re.search(r'[\u4e00-\u9fff]', text))


def preserve_technical_terms(text: str, target_lang: str) -> str:
    """确保专业术语在翻译后仍然保留原样"""
    for term in TECHNICAL_TERMS:
        pattern = re.compile(re.escape(term), re.IGNORECASE)
        text = pattern.sub(term, text)
    return text


def remove_bold_markers(text: str) -> str:
    """删除文本中的**粗体标记，同时清理Markdown语法"""
    if not text:
        return text
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    return text


def truncate_summary(text: str, max_len: int = 10) -> str:
    """截取摘要（用于标题）"""
    if not text:
        return "故障分析"
    text = remove_bold_markers(text)
    if len(text) <= max_len:
        return text
    for punct in ['。', '，', '、', '；', '：', '？', '！', '.', ',', ';', ':', '?', '!']:
        if punct in text[:max_len+5]:
            pos = text[:max_len+5].rfind(punct)
            if pos > 0:
                return text[:pos]
    return text[:max_len] + "…"


def get_llm_client():
    """获取LLM客户端"""
    if not DEEPSEEK_API_KEY:
        return None
    
    try:
        from openai import OpenAI
        return OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)
    except:
        return None


def call_llm(prompt: str, max_tokens: int = 4000, temperature: float = 0.3) -> str:
    """调用DeepSeek LLM"""
    client = get_llm_client()
    if not client:
        return "LLM配置错误"
    
    try:
        response = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            max_tokens=max_tokens
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"调用失败: {str(e)}"


def translate_to_en(text: str) -> str:
    """翻译为英文（保留专业术语）"""
    if not text or (not is_chinese(text) and not any(c.isalpha() for c in text)):
        return text
    
    client = get_llm_client()
    if not client:
        return text
    
    terms_str = ", ".join(TECHNICAL_TERMS)
    prompt = f"请将以下文本翻译成英文。注意：{terms_str} 等专业术语保持原样不翻译。\n\n文本：{text}\n\n只输出翻译结果，不要其他内容："
    
    try:
        response = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except:
        return text


def translate_to_zh(text: str) -> str:
    """翻译为中文（保留专业术语）"""
    if not text or (is_chinese(text) and not any(c.isalpha() for c in text)):
        return text
    
    client = get_llm_client()
    if not client:
        return text
    
    terms_str = ", ".join(TECHNICAL_TERMS)
    prompt = f"请将以下文本翻译成中文。注意：{terms_str} 等专业术语保持原样不翻译。\n\n文本：{text}\n\n只输出翻译结果，不要其他内容："
    
    try:
        response = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except:
        return text


def safe_json_parse(response: str, default_value: dict = None) -> dict:
    """安全解析JSON，支持降级处理"""
    if default_value is None:
        default_value = {}
    
    try:
        start = response.find('{')
        end = response.rfind('}') + 1
        if start != -1 and end > start:
            return json.loads(response[start:end])
        start = response.find('[')
        end = response.rfind(']') + 1
        if start != -1 and end > start:
            return json.loads(response[start:end])
    except json.JSONDecodeError:
        pass
    
    return default_value


# ==================== 联网搜索（双语双向，中文优先）====================

def web_search_dual(query: str, lang: str) -> str:
    """调用共享双语双向联网检索模块。"""
    return shared_web_search_dual(
        query=query,
        lang=lang,
        translate_to_en=translate_to_en,
        max_results_each=3,
        max_output=5,
    )


from knowledge_base_utils import SupabaseKnowledgeDB
from web_search_utils import web_search_dual as shared_web_search_dual


def create_supabase_knowledge_db() -> SupabaseKnowledgeDB:
    return SupabaseKnowledgeDB(
        SUPABASE_URL,
        SUPABASE_SERVICE_ROLE_KEY,
        translate_to_en=translate_to_en,
        translate_to_zh=translate_to_zh,
        ui_lang_getter=lambda: st.session_state.get("lang", "zh"),
    )


# ==================== 管理员弹窗 ====================

@st.dialog("管理员设置", width="large")
def admin_settings_dialog():
    """管理员设置弹窗"""
    lang = st.session_state.get("lang", "zh")
    
    if "admin_logged_in" not in st.session_state:
        st.session_state.admin_logged_in = False
    
    if not st.session_state.admin_logged_in:
        st.subheader(get_text("admin_login"))
        username = st.text_input(get_text("username"))
        password = st.text_input(get_text("password"), type="password")
        if st.button(get_text("login_btn")):
            if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
                st.session_state.admin_logged_in = True
                st.rerun()
            else:
                st.error("用户名或密码错误")
        return
    
    st.success(get_text("logged_in"))
    
    # 连接状态
    st.subheader(get_text("db_status"))
    
    col1, col2 = st.columns(2)
    with col1:
        if DEEPSEEK_API_KEY:
            st.success(f"✅ {get_text('llm_configured')}")
        else:
            st.error(f"❌ {get_text('llm_not_configured')}")
    
    with col2:
        if SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY:
            st.success(f"✅ {get_text('db_connected')}")
        else:
            st.error(f"❌ {get_text('db_disconnected')}")
    
    st.markdown("---")
    
    # 知识库管理
    st.subheader(get_text("knowledge_base_title"))
    
    if "knowledge_db" not in st.session_state:
        st.session_state.knowledge_db = create_supabase_knowledge_db()
    
    kb = st.session_state.knowledge_db
    categories = kb.categories
    
    selected_cat = st.selectbox(get_text("category"), categories)
    items = kb.get_knowledge(selected_cat, lang)
    
    st.write(f"共 {len(items)} 条记录")
    
    if items:
        with st.container(height=300):
            for idx, item in enumerate(items):
                col1, col2 = st.columns([10, 1])
                with col1:
                    display_item = remove_bold_markers(item[:150] + "..." if len(item) > 150 else item)
                    st.write(f"{idx+1}. {display_item}")
                with col2:
                    if st.button("❌", key=f"del_{selected_cat}_{idx}"):
                        kb.delete_knowledge(selected_cat, item)
                        st.rerun()
    else:
        st.info(get_text("no_entries"))
    
    new_item = st.text_area(get_text("entry_content"), height=80, 
                            placeholder=get_text("entry_placeholder"))
    if st.button(get_text("add_entry")):
        if new_item.strip():
            kb.add_knowledge(selected_cat, new_item.strip())
            st.rerun()
    
    st.markdown("---")
    
    # 导出/导入
    col_exp, col_imp = st.columns(2)
    
    with col_exp:
        st.markdown(f"**{get_text('export_kb')}**")
        if st.button("📥 " + get_text("export_kb"), key="export_btn"):
            df = kb.export_to_dataframe()
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name="KnowledgeBase", index=False)
            st.download_button(
                label="点击下载",
                data=output.getvalue(),
                file_name=f"knowledge_base_{datetime.now().strftime('%Y%m%d')}.xlsx",
                key="download_kb"
            )
    
    with col_imp:
        st.markdown(f"**{get_text('import_kb')}**")
        uploaded = st.file_uploader("选择Excel文件", type=["xlsx", "xls"], key="kb_upload")
        if uploaded:
            try:
                df = pd.read_excel(uploaded)
                count = kb.import_from_dataframe(df)
                st.success(get_text("import_success").format(count=count))
                st.rerun()
            except Exception as e:
                st.error(f"导入失败: {e}")


# ==================== 鱼骨图生成（放大版）====================

def setup_chinese_font():
    """设置中文字体（从网络加载）"""
    try:
        font_url = "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/SimplifiedChinese/NotoSansCJKsc-Regular.otf"
        font_path = "/tmp/NotoSansCJKsc-Regular.otf"
        
        import os
        if not os.path.exists(font_path):
            response = requests.get(font_url, timeout=30)
            with open(font_path, 'wb') as f:
                f.write(response.content)
        
        fontManager.addfont(font_path)
        plt.rcParams['font.sans-serif'] = ['Noto Sans CJK SC', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False
        return True
    except:
        plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
        return False


def create_fishbone_image(fishbone: FishboneAnalysis, lang: str = "zh") -> bytes:
    """生成鱼骨图图片（支持双语）"""
    setup_chinese_font()
    
    fig, ax = plt.subplots(figsize=(16, 10))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 12)
    ax.axis('off')
    
    # 主骨
    ax.plot([2, 14], [6, 6], 'k-', linewidth=3)
    
    # 箭头
    ax.annotate('', xy=(14, 6), xytext=(13.2, 6),
                arrowprops=dict(arrowstyle='->', lw=3, color='black'))
    
    # 鱼头标签
    ax.text(14.5, 6, get_text("symptom")[:25] if lang == "zh" else "Failure", 
            fontsize=12, va='center', fontweight='bold')
    
    cat_names_zh = {"Man": "人", "Machine": "机", "Material": "料",
                    "Method": "法", "Environment": "环", "Measurement": "测"}
    
    # 获取双语原因数据
    categories_data = [
        ("Man", fishbone.get_causes(lang, "Man"), 4, 9.5),
        ("Machine", fishbone.get_causes(lang, "Machine"), 4, 8.2),
        ("Material", fishbone.get_causes(lang, "Material"), 4, 6.9),
        ("Method", fishbone.get_causes(lang, "Method"), 5, 5.1),
        ("Environment", fishbone.get_causes(lang, "Environment"), 5, 3.8),
        ("Measurement", fishbone.get_causes(lang, "Measurement"), 5, 2.5)
    ]
    
    for cat_key, causes, spine_x, spine_y in categories_data:
        display_name = cat_names_zh.get(cat_key, cat_key) if lang == "zh" else cat_key
        
        # 脊骨
        ax.plot([2.5, spine_x], [6, spine_y], 'k-', linewidth=1.5)
        ax.text(spine_x - 0.5, spine_y, display_name, fontsize=11,
                ha='right', va='center', fontweight='bold')
        
        # 原因分支
        for j, cause in enumerate(causes[:4]):
            cause_text = remove_bold_markers(cause[:25] + "…" if len(cause) > 25 else cause)
            y_offset = spine_y + 0.6 + j * 0.45 if j % 2 == 0 else spine_y - 0.6 - (j-1) * 0.45
            ax.plot([spine_x, spine_x + 0.8], [spine_y, y_offset], 'k:', linewidth=0.8, alpha=0.6)
            ax.text(spine_x + 0.9, y_offset, cause_text, fontsize=8, va='center')
    
    ax.text(8, 11.2, remove_bold_markers(get_text("fishbone_title")), fontsize=16, ha='center', fontweight='bold')
    
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close()
    
    return buf.getvalue()


# ==================== 失效等级分类 ====================

STAGES = {
    0: {"zh": "正常", "en": "Normal"},
    1: {"zh": "轻微异常", "en": "Minor Anomaly"},
    2: {"zh": "中度异常", "en": "Moderate Anomaly"},
    3: {"zh": "严重故障", "en": "Critical Failure"}
}


def classify_stage(symptom: str) -> Tuple[int, float]:
    """分类故障等级"""
    symptom_lower = symptom.lower()
    
    keywords = {
        1: ["闪烁", "弱光", "色偏", "flicker", "dim", "intermittent"],
        2: ["烧焦", "膨胀", "变形", "burn", "swell", "deform", "黑化"],
        3: ["短路", "冒烟", "起火", "爆炸", "short", "smoke", "fire", "explode"]
    }
    
    scores = {0: 0, 1: 0, 2: 0, 3: 0}
    
    for stage, kw_list in keywords.items():
        for kw in kw_list:
            if kw in symptom_lower:
                scores[stage] += 1
    
    if max(scores.values()) == 0:
        return 1, 0.6
    
    best = max(scores, key=scores.get)
    confidence = min(0.95, 0.5 + scores[best] * 0.1)
    return best, confidence


def get_stage_name(stage: int, lang: str) -> str:
    return STAGES.get(stage, STAGES[1])[lang]


# ==================== 双语LLM分析（一次调用）====================

def call_bilingual_analysis(product_name: str, symptom_en: str, installation_en: str, 
                            temperature: str, context_info: str) -> dict:
    """
    一次LLM调用，输出双语分析结果
    
    返回格式：
    {
        "five_why": [...],
        "fishbone": {...},
        "root_cause_en": "...", "root_cause_zh": "...",
        "interim_actions_en": [...], "interim_actions_zh": [...],
        "permanent_actions_en": [...], "permanent_actions_zh": [...],
        "preventive_actions_en": [...], "preventive_actions_zh": [...]
    }
    """
    
    terms_str = ", ".join(TECHNICAL_TERMS)
    
    prompt = f"""You are a professional failure analysis engineer. Perform a thorough root cause analysis.

## Input Information
- Product: {product_name}
- Symptom: {symptom_en}
- Installation: {installation_en if installation_en else 'Not specified'}
- Temperature: {temperature if temperature else 'Not specified'}

## Context (Similar Cases & Industry Knowledge)
{context_info}

## Output Requirements

Generate a bilingual analysis (English and Chinese). **Critical Rules:**
1. Technical terms ({terms_str}) MUST remain in their original form in BOTH languages
2. All outputs must be in valid JSON format
3. Chinese content must be 100% Chinese characters (except technical terms listed above) - NO English sentences in Chinese fields
4. English content must be 100% English
5. For 5-Why: 5 levels, each with question (the "why" inquiry) and answer
6. For Fishbone: 3-5 causes per category
7. For actions: 2-4 interim, 3-5 permanent, 2-3 preventive

## Output JSON Structure:

{{
    "five_why": [
        {{"level": 1, "question_en": "Why...?", "question_zh": "为什么...？", 
          "answer_en": "Because...", "answer_zh": "因为...", "confidence": 0.95}},
        ... (5 levels)
    ],
    "fishbone": {{
        "man_en": ["cause1", "cause2"], "man_zh": ["原因1", "原因2"],
        "machine_en": [...], "machine_zh": [...],
        "material_en": [...], "material_zh": [...],
        "method_en": [...], "method_zh": [...],
        "environment_en": [...], "environment_zh": [...],
        "measurement_en": [...], "measurement_zh": [...]
    }},
    "root_cause_en": "The verified root cause in English...",
    "root_cause_zh": "验证后的根本原因中文...",
    "root_cause_confidence": 0.9,
    "interim_actions_en": ["ICA1", "ICA2"], "interim_actions_zh": ["临时措施1", "临时措施2"],
    "permanent_actions_en": ["PCA1", "PCA2", "PCA3"], "permanent_actions_zh": ["永久措施1", "永久措施2", "永久措施3"],
    "preventive_actions_en": ["PA1", "PA2"], "preventive_actions_zh": ["预防措施1", "预防措施2"]
}}

**IMPORTANT**: The Chinese fields (question_zh, answer_zh, man_zh, etc.) must contain ONLY Chinese characters and technical terms from the whitelist.

Generate only the JSON, no other text."""
    
    response = call_llm(prompt, max_tokens=5000, temperature=0.3)
    
    result = safe_json_parse(response)
    
    if not result:
        result = {}
    
    # 确保 five_why 有5层
    if "five_why" not in result or len(result.get("five_why", [])) != 5:
        result["five_why"] = []
        for i in range(1, 6):
            result["five_why"].append({
                "level": i,
                "question_en": f"Why did the failure occur? (Level {i})",
                "question_zh": f"为什么会发生这个故障？（第{i}层）",
                "answer_en": "Further analysis needed.",
                "answer_zh": "需要进一步分析。",
                "confidence": 0.6
            })
    
    # 确保 fishbone 所有维度都存在
    fishbone_default = {
        "man_en": [], "man_zh": [], "machine_en": [], "machine_zh": [],
        "material_en": [], "material_zh": [], "method_en": [], "method_zh": [],
        "environment_en": [], "environment_zh": [], "measurement_en": [], "measurement_zh": []
    }
    if "fishbone" not in result:
        result["fishbone"] = fishbone_default
    else:
        for key in fishbone_default:
            if key not in result["fishbone"]:
                result["fishbone"][key] = []
    
    # 确保 root_cause 字段存在
    if "root_cause_en" not in result:
        result["root_cause_en"] = "Analysis completed. Review 5-Why for details."
    if "root_cause_zh" not in result:
        result["root_cause_zh"] = "分析完成。详情请参考5-Why分析。"
    if "root_cause_confidence" not in result:
        result["root_cause_confidence"] = 0.7
    
    # 确保措施字段存在
    for key in ["interim_actions_en", "interim_actions_zh", 
                "permanent_actions_en", "permanent_actions_zh",
                "preventive_actions_en", "preventive_actions_zh"]:
        if key not in result:
            result[key] = ["Review analysis for recommendations"]
    
    return result


# ==================== 时序分析器 ====================

class TimeSeriesAnalyzer:
    """时序分析器 - SPC控制图"""
    
    @staticmethod
    def parse_data(data_text: str) -> Optional[pd.DataFrame]:
        try:
            df = pd.read_csv(io.StringIO(data_text))
            if 'production_qty' in df.columns:
                df['production_qty'] = pd.to_numeric(df['production_qty'], errors='coerce')
            if 'failure_qty' in df.columns:
                df['failure_qty'] = pd.to_numeric(df['failure_qty'], errors='coerce')
            return df.dropna()
        except:
            return None
    
    @staticmethod
    def parse_excel(file) -> Optional[pd.DataFrame]:
        try:
            if file.name.endswith('.csv'):
                df = pd.read_csv(file)
            else:
                df = pd.read_excel(file)
            if 'production_qty' in df.columns:
                df['production_qty'] = pd.to_numeric(df['production_qty'], errors='coerce')
            if 'failure_qty' in df.columns:
                df['failure_qty'] = pd.to_numeric(df['failure_qty'], errors='coerce')
            return df.dropna()
        except:
            return None
    
    @staticmethod
    def analyze_trend(df: pd.DataFrame) -> dict:
        if df is None or len(df) == 0:
            return {"has_data": False}
        
        total_failures = df['failure_qty'].sum()
        total_production = df['production_qty'].sum()
        overall_rate = total_failures / total_production * 100 if total_production > 0 else 0
        
        df['defect_rate'] = df['failure_qty'] / df['production_qty'] * 100
        recent_avg = df['defect_rate'].tail(3).mean() if len(df) >= 3 else df['defect_rate'].mean()
        
        return {
            "has_data": True,
            "overall_rate": overall_rate,
            "recent_rate": recent_avg,
            "trend": "上升" if recent_avg > overall_rate else "下降" if recent_avg < overall_rate else "稳定",
            "is_stable": True,
            "total_samples": len(df)
        }
    
    @staticmethod
    def create_spc_chart(df: pd.DataFrame) -> go.Figure:
        total_failures = df['failure_qty'].sum()
        total_production = df['production_qty'].sum()
        p_bar = total_failures / total_production if total_production > 0 else 0
        
        df['defect_rate'] = df['failure_qty'] / df['production_qty'] * 100
        
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=df.index,
            y=df['defect_rate'],
            mode='lines+markers',
            name='Defect Rate (%)'
        ))
        fig.add_hline(y=p_bar * 100, line_dash="dash", line_color="green",
                      annotation_text=f"Mean: {p_bar*100:.2f}%")
        
        fig.update_layout(
            title="SPC Control Chart",
            xaxis_title="Sample",
            yaxis_title="Defect Rate (%)",
            height=450
        )
        return fig


# ==================== 关联规则挖掘（双语）====================

def mine_association_rules_bilingual(symptom_en: str, installation_en: str, 
                                      temperature: str, lang: str) -> List[dict]:
    """挖掘关联规则（双语输出）"""
    
    output_lang = "Chinese" if lang == "zh" else "English"
    
    prompt = f"""Based on the failure information, discover potential association rules.

Symptom: {symptom_en}
Installation: {installation_en if installation_en else 'unknown'}
Temperature: {temperature if temperature else 'unknown'}

Output 2-3 association rules as JSON array. The explanation MUST be in {output_lang}.

Example for Chinese: {{"antecedents": ["高温>70°C", "上层安装"], "consequents": ["进水短路"], "confidence": 0.85, "explanation": "当安装在上层黑色金属表面的灯具遇到高温时，密封材料膨胀产生间隙，雨水渗入导致短路"}}

Example for English: {{"antecedents": ["high temperature >70°C", "upper installation"], "consequents": ["water ingress short circuit"], "confidence": 0.85, "explanation": "When luminaires installed on upper black metal surfaces experience high temperatures, seal materials expand creating gaps, allowing water ingress and short circuits"}}

Output only the JSON array, no other text."""
    
    response = call_llm(prompt, max_tokens=800, temperature=0.4)
    try:
        start = response.find('[')
        end = response.rfind(']') + 1
        if start != -1:
            rules = json.loads(response[start:end])
            for rule in rules[:3]:
                if "explanation" in rule:
                    rule["explanation"] = remove_bold_markers(rule["explanation"])
            return rules[:3]
    except:
        pass
    
    if lang == "zh":
        return [{
            "antecedents": ["分析中"],
            "consequents": ["待确认"],
            "confidence": 0.5,
            "explanation": "基于当前信息未发现明显的关联规则，建议收集更多数据。"
        }]
    else:
        return [{
            "antecedents": ["analyzing"],
            "consequents": ["to be confirmed"],
            "confidence": 0.5,
            "explanation": "No significant association rules found. More data collection recommended."
        }]


# ==================== 报告生成器（双语，无星号）====================

def generate_fa_report(result: FailureAnalysisResult, lang: str) -> str:
    """生成FA报告（纯指定语言，专业术语保留，无星号标记）"""
    stage_name = get_stage_name(result.failure_stage, lang)
    stage_emoji = {0: "✅", 1: "⚠️", 2: "🔥", 3: "🚨"}.get(result.failure_stage, "📌")
    
    if lang == "zh":
        symptom_text = result.symptom
        root_cause = remove_bold_markers(result.root_cause_zh)
        interim_actions = [remove_bold_markers(a) for a in result.interim_actions_zh]
        permanent_actions = [remove_bold_markers(a) for a in result.permanent_actions_zh]
        preventive_actions = [remove_bold_markers(a) for a in result.preventive_actions_zh]
    else:
        symptom_text = result.symptom_en
        root_cause = remove_bold_markers(result.root_cause_en)
        interim_actions = [remove_bold_markers(a) for a in result.interim_actions_en]
        permanent_actions = [remove_bold_markers(a) for a in result.permanent_actions_en]
        preventive_actions = [remove_bold_markers(a) for a in result.preventive_actions_en]
    
    fault_summary = symptom_text[:30] + "..." if len(symptom_text) > 30 else symptom_text
    
    info_table = f"""
## {get_text('report_info')}

| {get_text('project_name_header')} | {result.project_name if result.project_name else '-'} |
| {get_text('product_name_header')} | {result.product_name} |
| {get_text('fault_summary')} | {fault_summary} |
| {get_text('report_date')} | {datetime.now().strftime('%Y-%m-%d')} |
| {get_text('analyst')} | {result.analyst_name if result.analyst_name else '-'} {f"({result.analyst_title})" if result.analyst_title else ''} |

"""
    
        # 失效等级章节（包含等级定义表格）
    stage_section = f"""
## {get_text('stage_label')}

{stage_emoji} {stage_name} ({get_text('confidence')}: {result.root_cause_confidence:.0%})

### {get_text('stage_table_title')}

| {get_text('stage_table_header_level')} | {get_text('stage_table_header_name')} | {get_text('stage_table_header_characteristics')} | {get_text('stage_table_header_action')} |
|------|------|----------|----------|
| Stage 0 | {get_text('stage_0')} | {get_text('stage_0_char')} | {get_text('stage_0_action')} |
| Stage 1 | {get_text('stage_1')} | {get_text('stage_1_char')} | {get_text('stage_1_action')} |
| Stage 2 | {get_text('stage_2')} | {get_text('stage_2_char')} | {get_text('stage_2_action')} |
| Stage 3 | {get_text('stage_3')} | {get_text('stage_3_char')} | {get_text('stage_3_action')} |

"""
    
    five_why_table = f"| Level | {get_text('question_label')} | {get_text('answer_label')} | {get_text('confidence')} |\n|-------|----------|--------|------------|\n"
    five_why_list = ""
    
    for item in result.five_why:
        if lang == "zh":
            question = remove_bold_markers(item.question_zh)
            answer = remove_bold_markers(item.answer_zh)
        else:
            question = remove_bold_markers(item.question_en)
            answer = remove_bold_markers(item.answer_en)
        
        q_short = question[:45] + "..." if len(question) > 45 else question
        a_short = answer[:55] + "..." if len(answer) > 55 else answer
        five_why_table += f"| Why-{item.level} | {q_short} | {a_short} | {item.confidence:.0%} |\n"
        five_why_list += f"\n**Why-{item.level}**: {question}\n→ {answer}\n({get_text('confidence')}: {item.confidence:.0%})\n"
    
    five_why_section = f"""
## {get_text('five_why_title')}

### {get_text('table_format')}

{five_why_table}

### {get_text('detailed_description')}

{five_why_list}
"""
    
    root_cause_section = f"""
## {get_text('root_cause_title')}

{root_cause}

"""
    
    fishbone_dict = result.fishbone.to_dict(lang)
    fishbone_text = ""
    cat_names_zh = {"Man": "人", "Machine": "机", "Material": "料",
                    "Method": "法", "Environment": "环", "Measurement": "测"}
    
    for cat, causes in fishbone_dict.items():
        if causes:
            display_cat = cat_names_zh.get(cat, cat) if lang == "zh" else cat
            fishbone_text += f"\n{display_cat}:\n" + "\n".join([f"- {remove_bold_markers(c)}" for c in causes[:4]]) + "\n"
    
    fishbone_section = f"""
## {get_text('fishbone_title')}

{fishbone_text}
"""
    
    actions_text = f"""
### {get_text('interim_actions')}
{chr(10).join(f'{i+1}. {a}' for i, a in enumerate(interim_actions[:3]))}

### {get_text('permanent_actions')}
{chr(10).join(f'{i+1}. {a}' for i, a in enumerate(permanent_actions[:3]))}

### {get_text('preventive_actions')}
{chr(10).join(f'{i+1}. {a}' for i, a in enumerate(preventive_actions[:2]))}
"""
    
    return info_table + stage_section + five_why_section + root_cause_section + fishbone_section + actions_text


def generate_8d_report(result: FailureAnalysisResult, lang: str) -> str:
    """生成8D报告（纯指定语言，专业术语保留，无星号标记）"""
    stage_name = get_stage_name(result.failure_stage, lang)
    stage_emoji = {0: "✅", 1: "⚠️", 2: "🔥", 3: "🚨"}.get(result.failure_stage, "📌")
    
    if lang == "zh":
        symptom_text = remove_bold_markers(result.symptom)
        installation_text = remove_bold_markers(result.installation)
        root_cause = remove_bold_markers(result.root_cause_zh)
        interim_actions = [remove_bold_markers(a) for a in result.interim_actions_zh]
        permanent_actions = [remove_bold_markers(a) for a in result.permanent_actions_zh]
        preventive_actions = [remove_bold_markers(a) for a in result.preventive_actions_zh]
        
        five_why_items = []
        for item in result.five_why:
            five_why_items.append({
                "level": item.level,
                "question": remove_bold_markers(item.question_zh),
                "answer": remove_bold_markers(item.answer_zh)
            })
        
        fishbone_dict = result.fishbone.to_dict("zh")
        cat_names = {"Man": "人", "Machine": "机", "Material": "料",
                     "Method": "法", "Environment": "环", "Measurement": "测"}
        
        d1_table_data = {
            "headers": [get_text("team_role"), get_text("team_responsibility")],
            "rows": [
                [get_text("team_leader"), get_text("team_leader_resp")],
                [get_text("design_engineer"), get_text("design_engineer_resp")],
                [get_text("quality_engineer"), get_text("quality_engineer_resp")]
            ]
        }
        
        d6_table_data = {
            "headers": [get_text("verification_item"), get_text("verification_method"), get_text("verification_criteria")],
            "rows": [
                [get_text("function_test"), get_text("function_test_method"), get_text("function_test_criteria")],
                [get_text("durability_test"), get_text("durability_test_method"), get_text("durability_test_criteria")]
            ]
        }
        
        d8_items = [
            get_text("d8_recognition_1"),
            get_text("d8_recognition_2"),
            get_text("d8_recognition_3")
        ]
        
        d2_location_value = get_text("installation_site")
        d2_who_value = get_text("field_maintenance_team")
        d2_why_value = get_text("under_investigation")
        d2_how_value = get_text("failure_during_operation")
        d2_how_many_value = stage_name
        
    else:
        symptom_text = remove_bold_markers(result.symptom_en)
        installation_text = remove_bold_markers(result.installation_en)
        root_cause = remove_bold_markers(result.root_cause_en)
        interim_actions = [remove_bold_markers(a) for a in result.interim_actions_en]
        permanent_actions = [remove_bold_markers(a) for a in result.permanent_actions_en]
        preventive_actions = [remove_bold_markers(a) for a in result.preventive_actions_en]
        
        five_why_items = []
        for item in result.five_why:
            five_why_items.append({
                "level": item.level,
                "question": remove_bold_markers(item.question_en),
                "answer": remove_bold_markers(item.answer_en)
            })
        
        fishbone_dict = result.fishbone.to_dict("en")
        cat_names = {"Man": "Man", "Machine": "Machine", "Material": "Material",
                     "Method": "Method", "Environment": "Environment", "Measurement": "Measurement"}
        
        d1_table_data = {
            "headers": [get_text("team_role"), get_text("team_responsibility")],
            "rows": [
                [get_text("team_leader"), get_text("team_leader_resp")],
                [get_text("design_engineer"), get_text("design_engineer_resp")],
                [get_text("quality_engineer"), get_text("quality_engineer_resp")]
            ]
        }
        
        d6_table_data = {
            "headers": [get_text("verification_item"), get_text("verification_method"), get_text("verification_criteria")],
            "rows": [
                [get_text("function_test"), get_text("function_test_method"), get_text("function_test_criteria")],
                [get_text("durability_test"), get_text("durability_test_method"), get_text("durability_test_criteria")]
            ]
        }
        
        d8_items = [
            "Root cause identified and confirmed",
            "Improvement actions defined",
            "Lessons learned added to knowledge base"
        ]
        
        d2_location_value = "Installation site"
        d2_who_value = "Field maintenance team"
        d2_why_value = "Under investigation"
        d2_how_value = "Failure occurred during operation"
        d2_how_many_value = stage_name
    
    fault_summary = symptom_text[:30] + "..." if len(symptom_text) > 30 else symptom_text
    symptom_clean = symptom_text.replace('\n', ' ')
    
    info_table = f"""
## {get_text('report_info')}

| {get_text('project_name_header')} | {result.project_name if result.project_name else '-'} |
| {get_text('product_name_header')} | {result.product_name} |
| {get_text('fault_summary')} | {fault_summary} |
| {get_text('report_date')} | {datetime.now().strftime('%Y-%m-%d')} |
| {get_text('analyst')} | {result.analyst_name if result.analyst_name else '-'} {f"({result.analyst_title})" if result.analyst_title else ''} |

"""
    
    d1_table = "| " + " | ".join(d1_table_data["headers"]) + " |\n|" + "|".join(["------" for _ in d1_table_data["headers"]]) + "|\n"
    for row in d1_table_data["rows"]:
        d1_table += "| " + " | ".join(row) + " |\n"
    
    d1 = f"""
## D1: {get_text('establish_team')}

{d1_table}
"""
    
    d2 = f"""
## D2: {get_text('problem_description')}

| {get_text('what')} | {symptom_clean[:400]} |
| {get_text('where')} | {installation_text if installation_text else d2_location_value} |
| {get_text('when')} | {datetime.now().strftime('%Y-%m-%d')} |
| {get_text('who')} | {d2_who_value} |
| {get_text('why')} | {d2_why_value} |
| {get_text('how')} | {d2_how_value} |
| {get_text('how_many')} | {d2_how_many_value} |

"""
    
    d3 = f"""
## D3: {get_text('interim_actions')}

{chr(10).join(f'{i+1}. {a}' for i, a in enumerate(interim_actions[:3]))}

"""
    
    fishbone_text = ""
    for cat, causes in fishbone_dict.items():
        if causes:
            display_cat = cat_names.get(cat, cat)
            fishbone_text += f"\n**{display_cat}**: {', '.join([remove_bold_markers(c) for c in causes[:3]])}\n"
    
    five_why_list = ""
    for item in five_why_items:
        five_why_list += f"\n**Why-{item['level']}**: {item['question']}\n→ {item['answer']}\n"
    
    d4 = f"""
## D4: {get_text('root_cause_analysis')}

### {get_text('fishbone_analysis_title')}
{fishbone_text}

### {get_text('five_why_analysis_title')}
{five_why_list}

### {get_text('verified_root_cause_title')}
{root_cause}

"""
    
    d5 = f"""
## D5: {get_text('permanent_actions')}

{chr(10).join(f'{i+1}. {a}' for i, a in enumerate(permanent_actions[:3]))}

"""
    
    d6_table = "| " + " | ".join(d6_table_data["headers"]) + " |\n|" + "|".join(["------" for _ in d6_table_data["headers"]]) + "|\n"
    for row in d6_table_data["rows"]:
        d6_table += "| " + " | ".join(row) + " |\n"
    
    d6 = f"""
## D6: {get_text('effectiveness_verification')}

{d6_table}
"""
    
    d7 = f"""
## D7: {get_text('preventive_actions')}

{chr(10).join(f'{i+1}. {a}' for i, a in enumerate(preventive_actions[:2]))}

"""
    
    d8_items_formatted = chr(10).join(f'- {item}' for item in d8_items)
    d8 = f"""
## D8: {get_text('team_recognition')}

{d8_items_formatted}

"""
    
    return info_table + d1 + d2 + d3 + d4 + d5 + d6 + d7 + d8


def create_word_document(report_content: str, result: FailureAnalysisResult,
                         uploaded_images: List[bytes] = None,
                         fishbone_image: bytes = None,
                         lang: str = "zh") -> io.BytesIO:
    """创建Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        
        doc = Document()
        
        title_summary = truncate_summary(result.symptom if lang == "zh" else result.symptom_en, 12)
        title = doc.add_heading(f"{result.product_name} - {title_summary}", level=1)
        
        lines = report_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if line.startswith('# '):
                doc.add_heading(remove_bold_markers(line[2:]), level=1)
            elif line.startswith('## '):
                doc.add_heading(remove_bold_markers(line[3:]), level=2)
            elif line.startswith('### '):
                doc.add_heading(remove_bold_markers(line[4:]), level=3)
            elif line.startswith('|') and '|' in line[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                if len(table_lines) >= 2:
                    header_cells = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                    if '---' in table_lines[1]:
                        data_lines = table_lines[2:]
                    else:
                        data_lines = table_lines[1:]
                    
                    if header_cells and data_lines:
                        table = doc.add_table(rows=1+len(data_lines), cols=len(header_cells))
                        table.style = 'Table Grid'
                        for col, cell_text in enumerate(header_cells):
                            table.cell(0, col).text = remove_bold_markers(cell_text)
                        
                        for row_idx, data_line in enumerate(data_lines):
                            cells = [c.strip() for c in data_line.split('|')[1:-1]]
                            for col_idx, cell_text in enumerate(cells):
                                if col_idx < len(header_cells):
                                    table.cell(row_idx+1, col_idx).text = remove_bold_markers(cell_text)
                        doc.add_paragraph()
                continue
            elif line:
                doc.add_paragraph(remove_bold_markers(line))
            else:
                doc.add_paragraph()
            
            i += 1
        
        if fishbone_image:
            doc.add_page_break()
            doc.add_heading(remove_bold_markers(get_text("fishbone_title")), level=2)
            img_stream = io.BytesIO(fishbone_image)
            doc.add_picture(img_stream, width=Inches(12))
        
        if uploaded_images and len(uploaded_images) > 0:
            doc.add_page_break()
            doc.add_heading(remove_bold_markers(get_text("fault_photos")), level=2)
            for idx, img_bytes in enumerate(uploaded_images[:5]):
                try:
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(4))
                    doc.add_paragraph(f"Figure {idx+1}: Fault photo")
                except:
                    doc.add_paragraph(f"[Image {idx+1} could not be displayed]")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except ImportError:
        buffer = io.BytesIO()
        buffer.write(report_content.encode('utf-8'))
        buffer.seek(0)
        return buffer


# ==================== 主分析函数（双语架构）====================

def run_analysis(product_name: str, symptom: str, project_name: str,
                 installation: str, temperature: str, lang: str,
                 timeseries_df: pd.DataFrame = None,
                 enable_web: bool = True,
                 enable_rules: bool = True,
                 analyst_name: str = "", analyst_title: str = "") -> FailureAnalysisResult:
    """执行故障分析（双语架构）"""
    
    # 1. 翻译用户输入为英文（用于分析核心）
    symptom_en = translate_to_en(symptom) if is_chinese(symptom) else symptom
    product_name_en = translate_to_en(product_name) if is_chinese(product_name) else product_name
    installation_en = translate_to_en(installation) if installation and is_chinese(installation) else installation
    
    # 2. 双语检索（中文优先，英文补充）
    context_parts = []
    
    kb = create_supabase_knowledge_db()
    kb_all = kb.search_knowledge_full(symptom, limit=10)
    kb_results_zh = [r for r in kb_all if is_chinese(r)]
    kb_results_en = [r for r in kb_all if r not in kb_results_zh]
    
    if kb_results_zh:
        context_parts.append("【中文知识库案例】\n" + "\n".join(f"- {remove_bold_markers(r[:200])}" for r in kb_results_zh[:5]))
    
    if kb_results_en:
        context_parts.append("【English Knowledge Base】\n" + "\n".join(f"- {remove_bold_markers(r[:200])}" for r in kb_results_en[:5]))
    
    if enable_web:
        web_results = web_search_dual(symptom, lang)
        if web_results and "未找到" not in web_results:
            context_parts.append(web_results)
    
    context_info = "\n\n---\n\n".join(context_parts) if context_parts else "No similar cases found in knowledge base."
    
    # 3. 分类等级
    stage, _ = classify_stage(symptom)
    
    # 4. 一次LLM调用输出双语分析结果
    bilingual_result = call_bilingual_analysis(
        product_name_en, symptom_en, installation_en, temperature, context_info
    )
    
    # 5. 构建5-Why列表
    five_why_items = []
    for item_data in bilingual_result.get("five_why", []):
        five_why_items.append(FiveWhyItem(
            level=item_data.get("level", 1),
            question_en=remove_bold_markers(item_data.get("question_en", "")),
            question_zh=remove_bold_markers(item_data.get("question_zh", "")),
            answer_en=remove_bold_markers(item_data.get("answer_en", "")),
            answer_zh=remove_bold_markers(item_data.get("answer_zh", "")),
            confidence=item_data.get("confidence", 0.7),
            verification_method="建议通过测试验证"
        ))
    
    while len(five_why_items) < 5:
        level = len(five_why_items) + 1
        five_why_items.append(FiveWhyItem(
            level=level,
            question_en=f"Why did the failure occur? (Level {level})",
            question_zh=f"为什么会发生这个故障？（第{level}层）",
            answer_en="Further analysis needed.",
            answer_zh="需要进一步分析。",
            confidence=0.6,
            verification_method="建议通过测试验证"
        ))
    
    # 6. 构建鱼骨图
    fishbone_data = bilingual_result.get("fishbone", {})
    fishbone = FishboneAnalysis(
        man_en=[remove_bold_markers(c) for c in fishbone_data.get("man_en", [])],
        man_zh=[remove_bold_markers(c) for c in fishbone_data.get("man_zh", [])],
        machine_en=[remove_bold_markers(c) for c in fishbone_data.get("machine_en", [])],
        machine_zh=[remove_bold_markers(c) for c in fishbone_data.get("machine_zh", [])],
        material_en=[remove_bold_markers(c) for c in fishbone_data.get("material_en", [])],
        material_zh=[remove_bold_markers(c) for c in fishbone_data.get("material_zh", [])],
        method_en=[remove_bold_markers(c) for c in fishbone_data.get("method_en", [])],
        method_zh=[remove_bold_markers(c) for c in fishbone_data.get("method_zh", [])],
        environment_en=[remove_bold_markers(c) for c in fishbone_data.get("environment_en", [])],
        environment_zh=[remove_bold_markers(c) for c in fishbone_data.get("environment_zh", [])],
        measurement_en=[remove_bold_markers(c) for c in fishbone_data.get("measurement_en", [])],
        measurement_zh=[remove_bold_markers(c) for c in fishbone_data.get("measurement_zh", [])]
    )
    
    # 7. 生成鱼骨图图片
    fishbone_image = create_fishbone_image(fishbone, lang)
    
    # 8. SPC分析
    spc_analysis = None
    if timeseries_df is not None and len(timeseries_df) > 0:
        spc_analysis = TimeSeriesAnalyzer.analyze_trend(timeseries_df)
    
    # 9. 关联规则
    association_rules = []
    if enable_rules:
        association_rules = mine_association_rules_bilingual(symptom_en, installation_en, temperature, lang)
    
    # 10. 返回结果
    return FailureAnalysisResult(
        case_id=str(uuid.uuid4())[:8],
        project_name=project_name,
        product_name=product_name,
        symptom=symptom,
        symptom_en=symptom_en,
        installation=installation,
        installation_en=installation_en,
        temperature=temperature,
        failure_stage=stage,
        five_why=five_why_items,
        fishbone=fishbone,
        root_cause_en=remove_bold_markers(bilingual_result.get("root_cause_en", "")),
        root_cause_zh=remove_bold_markers(bilingual_result.get("root_cause_zh", "")),
        root_cause_confidence=bilingual_result.get("root_cause_confidence", 0.7),
        interim_actions_en=[remove_bold_markers(a) for a in bilingual_result.get("interim_actions_en", [])],
        interim_actions_zh=[remove_bold_markers(a) for a in bilingual_result.get("interim_actions_zh", [])],
        permanent_actions_en=[remove_bold_markers(a) for a in bilingual_result.get("permanent_actions_en", [])],
        permanent_actions_zh=[remove_bold_markers(a) for a in bilingual_result.get("permanent_actions_zh", [])],
        preventive_actions_en=[remove_bold_markers(a) for a in bilingual_result.get("preventive_actions_en", [])],
        preventive_actions_zh=[remove_bold_markers(a) for a in bilingual_result.get("preventive_actions_zh", [])],
        internal_cases_used=len(kb_results_zh) + len(kb_results_en),
        external_sources_used=1 if enable_web and context_info != "No similar cases found in knowledge base." else 0,
        spc_analysis=spc_analysis,
        association_rules=association_rules,
        analyst_name=analyst_name,
        analyst_title=analyst_title,
        fishbone_image=fishbone_image
    )


# ==================== 主页面 ====================

def main():
    """主应用入口"""
    
    # 初始化 session state
    if "result" not in st.session_state:
        st.session_state.result = None
    if "current_report" not in st.session_state:
        st.session_state.current_report = None
    if "report_type" not in st.session_state:
        st.session_state.report_type = "fa"
    if "analysis_completed" not in st.session_state:
        st.session_state.analysis_completed = False  # 分析是否完成
    if "analyst_name" not in st.session_state:
        st.session_state.analyst_name = ""
    if "analyst_title" not in st.session_state:
        st.session_state.analyst_title = ""
    if "project_name" not in st.session_state:
        st.session_state.project_name = ""
    if "uploaded_images" not in st.session_state:
        st.session_state.uploaded_images = []
    
    # ==================== 侧边栏（显示用户信息和剩余次数）====================
    render_sidebar_user_info()
    
    with st.sidebar:
        st.markdown(f"### {get_text('sidebar_about')}")
        st.markdown(get_text("sidebar_principle"))
        st.markdown(get_text("sidebar_usage"))
        st.markdown("---")
        
        st.session_state.analyst_name = st.text_input(get_text("analyst_name"), placeholder=get_text("analyst_name_ph"))
        st.session_state.analyst_title = st.text_input(get_text("analyst_title"), placeholder=get_text("analyst_title_ph"))
        st.session_state.project_name = st.text_input(get_text("project_name_label"), placeholder=get_text("project_name_ph"))
        
        if st.session_state.analyst_name:
            st.success(f"{get_text('analyst')}: {st.session_state.analyst_name}")
        
        st.markdown("---")
        st.markdown(f"**{get_text('db_status')}**")
        if SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY:
            st.success(f"✅ {get_text('db_connected')}")
        else:
            st.error(f"❌ {get_text('db_disconnected')}")
        
        st.markdown("---")
        st.markdown(f"### {get_text('contact')}")
        st.markdown(get_text("contact_email"))
    
    # 右上角语言切换和齿轮
    col1, col2, col3, col4, col5 = st.columns([2, 2, 1, 1, 1])
    with col3:
        if st.button(get_text("lang_zh"), key="zh_btn"):
            set_app_language("zh")
            st.rerun()
    with col4:
        if st.button(get_text("lang_en"), key="en_btn"):
            set_app_language("en")
            st.rerun()
    with col5:
        if st.button("⚙️", key="settings_btn"):
            admin_settings_dialog()
    
    st.title(get_text("app_title"))
    st.caption(get_text("app_subtitle"))
    
    if not DEEPSEEK_API_KEY:
        st.error(get_text("api_error"))
        return
    
    # ==================== 检查剩余次数 ====================
    remaining, tier, expires_at, error = get_user_remaining_trials(st.session_state.user_id)
    
    if error:
        st.error(error)
        return
    
    has_trials = (tier == "pro") or (remaining > 0)
    
    if not has_trials:
        st.error(f"⚠️ {get_text('trials_insufficient')}")
        # 显示升级提示
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.info("💎 " + get_text("click_to_upgrade") + " - 请联系管理员或返回门户升级")
        return
    
    # ==================== 主表单 ====================
    st.markdown(f"### {get_text('basic_info')}")
    
    product_name = st.text_input(get_text("product_name"), placeholder=get_text("product_name_ph"))
    symptom = st.text_area(get_text("symptom"), height=120, placeholder=get_text("symptom_ph"))
    installation = st.text_input(get_text("installation"), placeholder=get_text("installation_ph"))
    
    col_date, col_batch = st.columns(2)
    with col_date:
        st.date_input(get_text("failure_date"), value=datetime.now().date())
    with col_batch:
        st.text_input(get_text("batch_no"), placeholder="LOT2024-001")
    
    temperature = st.text_input(get_text("site_temp"), placeholder=get_text("site_temp_ph"))
    
    st.markdown("---")
    st.markdown(f"### {get_text('image_section')}")
    
    uploaded_files = st.file_uploader(get_text("image_upload_hint"), type=["jpg", "jpeg", "png"], accept_multiple_files=True)
    if uploaded_files:
        st.session_state.uploaded_images = [img.getvalue() for img in uploaded_files]
        cols = st.columns(min(3, len(uploaded_files)))
        for idx, img in enumerate(uploaded_files[:3]):
            with cols[idx]:
                st.image(img, use_container_width=True)
    else:
        st.caption(get_text("no_images"))
    
    st.markdown("---")
    st.markdown(f"### {get_text('timeseries_section')}")
    
    enable_timeseries = st.checkbox(get_text("timeseries_checkbox"), value=False)
    timeseries_df = None
    
    if enable_timeseries:
        input_method = st.radio(get_text("timeseries_input_method"), [get_text("timeseries_paste"), get_text("timeseries_upload")], horizontal=True)
        
        if input_method == get_text("timeseries_paste"):
            paste_data = st.text_area(get_text("timeseries_paste_placeholder"), height=150)
            if paste_data:
                timeseries_df = TimeSeriesAnalyzer.parse_data(paste_data)
        else:
            timeseries_file = st.file_uploader(get_text("timeseries_file_hint"), type=["xlsx", "xls", "csv"])
            if timeseries_file:
                timeseries_df = TimeSeriesAnalyzer.parse_excel(timeseries_file)
                if timeseries_df is not None:
                    st.dataframe(timeseries_df.head(), use_container_width=True)
        
        template_df = pd.DataFrame({"date": ["2024-01-01", "2024-02-01"], "production_qty": [1000, 1100], "failure_qty": [5, 8]})
        st.download_button(get_text("download_template"), template_df.to_csv(index=False).encode(), "template.csv")
    
    st.markdown("---")
    st.markdown(f"### {get_text('advanced_options')}")
    
    col_adv1, col_adv2, col_adv3, col_adv4 = st.columns(4)
    with col_adv1:
        enable_web = st.checkbox(get_text("web_search"), value=True)
    with col_adv2:
        enable_rules = st.checkbox(get_text("rule_mining"), value=True)
    with col_adv3:
        enable_spc = st.checkbox(get_text("spc"), value=True)
    with col_adv4:
        enable_8d = st.checkbox(get_text("gen_8d"), value=True)
    
    st.markdown("---")
    
    # ==================== 分析按钮（扣费）====================
    if st.button(get_text("analyze_btn"), type="primary", use_container_width=True):
        if not product_name or not symptom:
            st.error(get_text("fill_required"))
        else:
            # 再次检查剩余次数
            remaining_check, tier_check, _, _ = get_user_remaining_trials(st.session_state.user_id)
            if tier_check != "pro" and remaining_check <= 0:
                st.error(get_text("trials_insufficient"))
            else:
                # 消耗次数
                success, new_remaining, error_msg = consume_trial(
                    st.session_state.user_id, 
                    "AI-FA", 
                    "深度故障分析"
                )
                
                if not success:
                    st.error(error_msg)
                else:
                    with st.spinner(get_text("analyzing")):
                        try:
                            result = run_analysis(
                                product_name=product_name,
                                symptom=symptom,
                                project_name=st.session_state.project_name,
                                installation=installation,
                                temperature=temperature,
                                lang=st.session_state.lang,
                                timeseries_df=timeseries_df if enable_timeseries else None,
                                enable_web=enable_web,
                                enable_rules=enable_rules,
                                analyst_name=st.session_state.analyst_name,
                                analyst_title=st.session_state.analyst_title
                            )
                            st.session_state.result = result
                            st.session_state.current_report = None
                            st.session_state.analysis_completed = True
                            st.success(get_text("success"))
                            st.rerun()
                        except Exception as e:
                            st.error(f"{get_text('error')}: {str(e)}")
    
    # ==================== 显示结果 ====================
    if st.session_state.result and st.session_state.analysis_completed:
        result = st.session_state.result
        lang = st.session_state.lang
        stage_name = get_stage_name(result.failure_stage, lang)
        
        st.markdown("---")
        st.info(f"**{get_text('stage_label')}**: {stage_name} ({get_text('confidence')}: {result.root_cause_confidence:.0%})")
        
        with st.expander(get_text("five_why_title"), expanded=True):
            for item in result.five_why:
                if lang == "zh":
                    question = remove_bold_markers(item.question_zh)
                    answer = remove_bold_markers(item.answer_zh)
                else:
                    question = remove_bold_markers(item.question_en)
                    answer = remove_bold_markers(item.answer_en)
                st.markdown(f"**Why-{item.level}**: {question}")
                st.markdown(f"→ {answer}")
                st.progress(item.confidence, text=f"{get_text('confidence')}: {item.confidence:.0%}")
                st.divider()
        
        with st.expander(get_text("fishbone_title")):
            if result.fishbone_image:
                st.image(result.fishbone_image, use_container_width=True)
            else:
                fishbone_dict = result.fishbone.to_dict(lang)
                cat_names_zh = {"Man": "人", "Machine": "机", "Material": "料",
                                "Method": "法", "Environment": "环", "Measurement": "测"}
                for cat, causes in fishbone_dict.items():
                    if causes:
                        display_cat = cat_names_zh.get(cat, cat) if lang == "zh" else cat
                        st.markdown(f"**{display_cat}**")
                        for c in causes[:4]:
                            st.markdown(f"- {remove_bold_markers(c)}")
        
        if result.spc_analysis and result.spc_analysis.get("has_data"):
            with st.expander(get_text("spc_title")):
                st.metric("Overall Rate", f"{result.spc_analysis['overall_rate']:.2f}%")
                st.metric("Trend", result.spc_analysis['trend'])
        
        if result.association_rules:
            with st.expander(get_text("rules_title")):
                for rule in result.association_rules:
                    antecedents = " + ".join(rule.get("antecedents", []))
                    consequents = " + ".join(rule.get("consequents", []))
                    explanation = remove_bold_markers(rule.get("explanation", ""))
                    st.info(f"{antecedents} → {consequents} (Confidence: {rule.get('confidence', 0):.0%})\n\n{explanation}")
        
        st.markdown(f"### {get_text('root_cause_title')}")
        if lang == "zh":
            st.success(remove_bold_markers(result.root_cause_zh))
        else:
            st.success(remove_bold_markers(result.root_cause_en))
        
        # ==================== 报告生成按钮（各自扣费）====================
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        
        with col_btn1:
            # FA 报告按钮
            if st.button(get_text("generate_fa_btn"), use_container_width=True):
                # 检查剩余次数
                remaining_check, tier_check, _, _ = get_user_remaining_trials(st.session_state.user_id)
                if tier_check != "pro" and remaining_check <= 0:
                    st.error(get_text("trials_insufficient"))
                else:
                    # 消耗次数
                    success, new_remaining, error_msg = consume_trial(
                        st.session_state.user_id, 
                        "AI-FA", 
                        "生成FA报告"
                    )
                    if not success:
                        st.error(error_msg)
                    else:
                        report = generate_fa_report(result, lang)
                        st.session_state.current_report = report
                        st.session_state.report_type = "FA"
                        st.rerun()
        
        with col_btn2:
            # 8D 报告按钮
            if st.button(get_text("generate_8d_btn"), use_container_width=True):
                remaining_check, tier_check, _, _ = get_user_remaining_trials(st.session_state.user_id)
                if tier_check != "pro" and remaining_check <= 0:
                    st.error(get_text("trials_insufficient"))
                else:
                    success, new_remaining, error_msg = consume_trial(
                        st.session_state.user_id, 
                        "AI-FA", 
                        "生成8D报告"
                    )
                    if not success:
                        st.error(error_msg)
                    else:
                        report = generate_8d_report(result, lang)
                        st.session_state.current_report = report
                        st.session_state.report_type = "8D"
                        st.rerun()
        
        with col_btn3:
            # 清除结果按钮（不扣费）
            if st.button(get_text("clear_btn"), use_container_width=True):
                st.session_state.result = None
                st.session_state.current_report = None
                st.session_state.analysis_completed = False
                st.rerun()
    
    # ==================== 显示报告预览 ====================
    if st.session_state.current_report and st.session_state.result:
        st.markdown("---")
        st.markdown(f"### {get_text('report_preview')}")
        with st.container(height=500):
            st.markdown(st.session_state.current_report)
        
        result = st.session_state.result
        lang = st.session_state.lang
        title_summary = truncate_summary(result.symptom if lang == "zh" else result.symptom_en, 12)
        filename = f"{result.product_name}_{title_summary}_{st.session_state.report_type}_Report_{datetime.now().strftime('%Y%m%d')}.docx"
        filename = re.sub(r'[\\/*?:"<>|]', '', filename)
        
        word_buffer = create_word_document(
            st.session_state.current_report,
            st.session_state.result,
            st.session_state.uploaded_images,
            st.session_state.result.fishbone_image,
            lang
        )
        
        st.download_button(
            label=get_text("download_word"),
            data=word_buffer,
            file_name=filename,
            use_container_width=True
        )


if __name__ == "__main__":
    setup_chinese_font()
    main()
